"""ApexNodes social-media media tools.

平台工具网关 P1(DESKTOP-CLOUD-CAPABILITY-PARITY-PD D1/D2):本插件自
hermes-cloud ``app/runtime_plugins/apexnodes-douyin-tools`` 迁入 fork,
桌面与云端同一份代码。默认走平台工具网关(``plugins/apexnodes_gateway.py``,
vendor key 永不出云);``TOOLS_GATEWAY_DISABLED=1`` 时回退迁移前的 master
内网端点路径(云端 P1 回退通道,P2 删)。

hc-254: tool names are platform-neutral (``social_download`` /
``media_transcribe``); the package directory keeps its
``apexnodes-douyin-tools`` name for plugin-registry stability.
"""

from __future__ import annotations

import json
import os
import re
import urllib.error
import urllib.request
import uuid
from pathlib import Path
from typing import Any

from tools.registry import tool_error, tool_result

try:  # 网关客户端(fork 内共用);导入失败时插件退回 legacy 路径,绝不拖垮装载。
    from plugins import apexnodes_gateway as _gateway
except Exception:  # pragma: no cover - 仅在裸拷贝部署等异常形态出现
    _gateway = None  # type: ignore[assignment]


def _use_gateway() -> bool:
    return _gateway is not None and _gateway.use_gateway()


# ── legacy master 直连路径(迁移前行为,原样保留;TOOLS_GATEWAY_DISABLED=1 或无网关配置时走) ──

def _api_base() -> str:
    base = (
        os.getenv("HERMES_PLATFORM_API_BASE")
        or os.getenv("HERMES_MASTER_API_BASE")
        or os.getenv("HERMES_SCHEDULER_API_BASE")
        or "http://host.docker.internal:8000/api/v1"
    )
    return base.rstrip("/")


def _agent_api_key() -> str:
    return (os.getenv("API_SERVER_KEY") or os.getenv("MODEL_API_KEY") or "").strip()


# hc-340: media_transcribe holds the master connection for the whole ASR poll,
# whose budget scales with audio length and is capped at 2h (hc-323). The old
# shared 600s (10min) cap was far below that, so any long video timed out at the
# plugin layer and forced a wasteful re-download + re-transcribe retry (and
# scared the user with a failure mid-way). Align the transcribe timeout with the
# master budget + download/ffmpeg headroom; light calls keep the short default.
_MEDIA_TRANSCRIBE_TIMEOUT_SECONDS = 2 * 60 * 60 + 10 * 60  # 2h10m


def _post(path: str, payload: dict[str, Any], *, timeout: int = 600) -> dict[str, Any]:
    api_key = _agent_api_key()
    if not api_key:
        raise RuntimeError("Agent API key is missing")
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    request = urllib.request.Request(
        f"{_api_base()}{path}",
        data=body,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            return json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        raw = exc.read().decode("utf-8", errors="replace")
        try:
            detail = json.loads(raw).get("detail", raw)
        except json.JSONDecodeError:
            detail = raw
        if isinstance(detail, dict):
            message = detail.get("message") or detail.get("code") or raw
            # hc-254: surface the machine-readable honesty signal so the model
            # offers the honest options instead of self-installing whisper.
            if detail.get("terminal_fallback_allowed") is False:
                message = (
                    f"{message}\n[terminal_fallback_allowed=false] 平台转写不可用，"
                    "请按失败选项处理（稍后重试 / 发视频文件 / 贴文字稿），"
                    "不要尝试自行安装或运行本地转写工具（如 whisper）。"
                )
        else:
            message = str(detail or raw)
        raise RuntimeError(message) from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"无法连接媒体服务: {exc.reason}") from exc


def _get(path: str, *, timeout: int = 60) -> dict[str, Any]:
    api_key = _agent_api_key()
    if not api_key:
        raise RuntimeError("Agent API key is missing")
    request = urllib.request.Request(
        f"{_api_base()}{path}",
        headers={"Authorization": f"Bearer {api_key}"},
        method="GET",
    )
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            return json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        raw = exc.read().decode("utf-8", errors="replace")
        try:
            detail = json.loads(raw).get("detail", raw)
        except json.JSONDecodeError:
            detail = raw
        message = detail.get("message") if isinstance(detail, dict) else str(detail or raw)
        raise RuntimeError(message or raw) from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"无法连接媒体服务: {exc.reason}") from exc


def _legacy_check() -> bool:
    api_key = _agent_api_key()
    if not api_key:
        return False
    request = urllib.request.Request(
        f"{_api_base()}/media/douyin/check",
        headers={"Authorization": f"Bearer {api_key}"},
        method="GET",
    )
    try:
        with urllib.request.urlopen(request, timeout=5) as response:
            data = json.loads(response.read().decode("utf-8"))
            return bool(data.get("tikhub_configured") and data.get("doubao_asr_configured") and data.get("ffmpeg_available"))
    except Exception:
        return False


def _check() -> bool:
    # 网关模式:可用性=配置齐(base+key);额度/上游故障在调用时以显式文案暴露,
    # 不在注册期做网络探测(桌面离线启动不应拖慢/误杀工具注册)。
    if _use_gateway():
        return bool(_gateway.agent_api_key())
    return _legacy_check()


SOCIAL_DOWNLOAD_SCHEMA = {
    "name": "social_download",
    "description": (
        "下载社媒视频（抖音/TikTok、小红书、快手、B站）到当前工作区的媒体空间。用户发来抖音/小红书/快手/B站分享链接/分享口令时，"
        "第一步必须优先使用本工具；不要用浏览器打开这些链接（反爬环境又慢又不可靠）。返回 video_path 和元数据。"
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "url": {
                "type": "string",
                "description": (
                    "直接传用户发来的原始分享文本即可（整段口令，含前后噪音文字也可以）；"
                    "平台会自动识别抖音（v.douyin.com）、小红书（xhslink.com / xiaohongshu.com）、"
                    "快手（v.kuaishou.com）、B站（bilibili.com / b23.tv）链接。"
                ),
            }
        },
        "required": ["url"],
    },
}

MEDIA_TRANSCRIBE_SCHEMA = {
    "name": "media_transcribe",
    "description": (
        "转写已下载的 video_path 或社媒视频链接 url（抖音/小红书/快手/B站，可直接传用户的原始分享文本，一步下载并转写）。"
        "用户要视频文案/逐字稿/拆解脚本时用本工具，禁止用浏览器或其他平台替代，也不要自行安装或运行本地转写工具（如 whisper）。"
        "返回完整转写和 transcript_path。"
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "video_path": {"type": "string", "description": "Video path returned by social_download."},
            "url": {"type": "string", "description": "Optional social video share text/URL; downloads and transcribes in one call."},
        },
    },
}


IMAGE_OCR_SCHEMA = {
    "name": "image_ocr",
    "description": (
        "识别社媒图文/图片里的文字（抖音/TikTok 图文、小红书图文、快手图片）。"
        "用户发来图文笔记/图片链接、或想「提取图片文字/读菜单价格/读截图文字」时用本工具：传社媒分享链接 url，"
        "或直接传图片地址 image_urls 列表。不要用浏览器打开这些链接。本工具只读字，不理解非文字画面/场景/构图。"
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "url": {
                "type": "string",
                "description": "社媒图文分享链接/分享口令（抖音图文、小红书图文、快手图片、TikTok photo），平台自动识别并提取图片。",
            },
            "image_urls": {
                "type": "array",
                "items": {"type": "string"},
                "description": "可选：直接给出的图片地址列表（已有图片 URL 时用，无需再传 url）。",
            },
            "prompt": {
                "type": "string",
                "description": "可选：用户对这些图片文字的具体需求（如「提取菜单价格」「读出截图里的文案」）。",
            },
        },
    },
}


def _legacy_check_image() -> bool:
    # hc-320: image OCR needs local RapidOCR; surfaced by /media/douyin/check.
    api_key = _agent_api_key()
    if not api_key:
        return False
    request = urllib.request.Request(
        f"{_api_base()}/media/douyin/check",
        headers={"Authorization": f"Bearer {api_key}"},
        method="GET",
    )
    try:
        with urllib.request.urlopen(request, timeout=5) as response:
            data = json.loads(response.read().decode("utf-8"))
            return bool(data.get("image_ocr_available"))
    except Exception:
        return False


def _check_image() -> bool:
    if _use_gateway():
        return bool(_gateway.agent_api_key())
    return _legacy_check_image()


# ── 网关模式的媒体编排(PD §8:网关只出直链/元数据,媒体本端自取;ASR 上传仅音轨) ──

def _gateway_resolve_download(url: str, *, timeout: float = 600) -> tuple[dict[str, Any], str]:
    """经网关解析分享链接:返回 (含元数据的结果 dict, 直链 media_url——可为空)。"""
    platform = _gateway.detect_platform(url)
    if not platform:
        raise _gateway.GatewayError(
            "无法从链接识别社媒平台（支持 抖音/小红书/快手/B站/TikTok/YouTube/Instagram/Twitter/公众号/视频号），"
            "请发送原始分享链接或分享口令。"
        )
    response = _gateway.request_json(
        "POST", f"/tools/v1/social/{platform}/download", {"url": url}, timeout=timeout
    )
    result = _gateway.unwrap(response)
    media_url = str(
        result.get("download_url") or result.get("media_url") or result.get("video_url") or ""
    ).strip()
    result.setdefault("platform", platform)
    return result, media_url


def _gateway_social_download(url: str) -> str:
    try:
        result, media_url = _gateway_resolve_download(url)
        if media_url:
            video_path = _gateway.download_media(
                media_url,
                headers=result.pop("download_headers", None),
                filename_hint=str(result.get("title") or ""),
            )
            result["video_path"] = str(video_path)
        else:
            result.pop("download_headers", None)
    except _gateway.GatewayError as exc:
        return tool_error(f"视频下载失败: {exc}")
    result.setdefault("ok", True)
    return tool_result(**result)


def _gateway_media_transcribe(video_path: str | None, url: str | None) -> str:
    meta: dict[str, Any] = {}
    local_path: Path | None = None
    media_url = ""
    try:
        if video_path:
            local_path = Path(video_path)
            if not local_path.exists():
                return tool_error(f"视频转写失败: 找不到视频文件 {video_path}，请先用 social_download 下载。")
        elif url:
            meta, media_url = _gateway_resolve_download(url)
            if media_url:
                try:
                    local_path = _gateway.download_media(
                        media_url,
                        headers=meta.pop("download_headers", None),
                        filename_hint=str(meta.get("title") or ""),
                    )
                except _gateway.GatewayError:
                    # 直链本地下载失败 → 退回契约基本形态:让网关自取媒体转写。
                    local_path = None

        if local_path is not None:
            upload_path = _gateway.extract_audio_for_asr(local_path) or local_path
            with open(upload_path, "rb") as fh:
                files = {
                    "file": (upload_path.name, fh, _gateway.guess_media_content_type(upload_path)),
                }
                response = _gateway.request_json(
                    "POST",
                    "/tools/v1/asr/transcribe",
                    files=files,
                    timeout=_MEDIA_TRANSCRIBE_TIMEOUT_SECONDS,
                )
        elif media_url:
            response = _gateway.request_json(
                "POST",
                "/tools/v1/asr/transcribe",
                {"media_url": media_url},
                timeout=_MEDIA_TRANSCRIBE_TIMEOUT_SECONDS,
            )
        else:
            return tool_error(
                "视频转写失败: 平台未能给出该内容的可下载媒体（可能是图文/直播链接），"
                "请确认是视频分享链接，或先用 social_download 检查。"
            )
    except _gateway.GatewayError as exc:
        return tool_error(f"视频转写失败: {exc}")

    transcript = str(response.get("text") or "").strip()
    title = str(meta.get("title") or "").strip()
    result: dict[str, Any] = {
        "ok": True,
        "transcript": transcript,
        "video_path": str(local_path) if local_path is not None else None,
        "title": title or None,
        "author": meta.get("author"),
        "audio_duration_seconds": response.get("duration_seconds"),
    }
    if response.get("cost_cents") is not None:
        result["cost_cents"] = response.get("cost_cents")
    transcript_path = _write_transcript_file(transcript, title)
    if transcript_path:
        result["transcript_path"] = transcript_path
    # hc-341: pre-render the verbatim transcript so "完整文案/逐字稿/下载文案" is delivered
    # as the original speech, copied verbatim — the agent only emits the MEDIA line.
    if transcript and _doc_file_delivery_available():
        doc_path = render_verbatim_transcript_docx(transcript, title)
        if doc_path:
            result["transcript_doc_path"] = doc_path
            result["transcript_doc_media_tag"] = f"MEDIA:{doc_path}"
            result["transcript_doc_instruction"] = _TRANSCRIPT_DOC_INSTRUCTION
    return tool_result(**result)


def _write_transcript_file(transcript: str, title: str) -> str | None:
    """把转写全文落成本地 txt(保持迁移前 transcript_path 的结果面)。"""
    if not transcript:
        return None
    try:
        dest = _document_cache_dir() / f"{_safe_filename(title or '视频转写')}_转写_{uuid.uuid4().hex[:8]}.txt"
        dest.write_text(transcript, encoding="utf-8")
        return str(dest)
    except Exception:  # noqa: BLE001 — 落盘失败不应吞掉转写本身
        return None


def _gateway_image_ocr(payload: dict[str, Any]) -> str:
    # image-ocr 的 OCR 在云侧执行;platform 路径段只作路由/白名单——裸图片 URL
    # 识别不出平台时回退 "douyin"(见 plugins/apexnodes_gateway.py 契约表)。
    probe_text = " ".join(
        [str(payload.get("url") or "")] + [str(u) for u in (payload.get("image_urls") or [])]
    )
    platform = _gateway.detect_platform(probe_text) or "douyin"
    try:
        response = _gateway.request_json(
            "POST", f"/tools/v1/social/{platform}/image-ocr", payload, timeout=600
        )
    except _gateway.GatewayError as exc:
        return tool_error(f"图片 OCR 失败: {exc}")
    return tool_result(**_gateway.unwrap(response))


def _routed_intent(args: dict) -> str | None:
    """hc-450: log-only marker the intent-router hook stamps onto a routed call;
    forwarded so the master can emit its structured intent_router_hit line."""
    value = str(args.get("routed_intent") or "").strip()
    return value[:64] or None


def _handle_social_download(args: dict, **_kwargs) -> str:
    if not isinstance(args, dict):
        return tool_error("social_download expects a JSON object argument")
    url = str(args.get("url") or args.get("share_url") or "").strip()
    if not url:
        return tool_error("请提供视频分享链接")
    if _use_gateway():
        return _gateway_social_download(url)
    payload: dict[str, Any] = {"url": url}
    if _routed_intent(args):
        payload["routed_intent"] = _routed_intent(args)
    try:
        return tool_result(**_post("/media/social-download", payload))
    except RuntimeError as exc:
        return tool_error(f"视频下载失败: {exc}")


# ── hc-341: verbatim transcript → .docx (programmatic, container-side) ────────
# A "完整文案 / 逐字稿 / 下载文案" request must deliver the ASR text copied verbatim.
# Prompt-only carve-outs didn't hold — the model kept restructuring the body when it
# round-tripped through doc_file_write. So we render the .docx here, straight from the
# transcribe result, and hand the agent a ready MEDIA tag: it only triggers the export,
# it never regenerates the text. Gated to the non-Feishu IM entries where a .docx is the
# delivery (Feishu keeps feishu_doc_write); mirrors apexnodes-doc-file-write's renderer.
_NON_FEISHU_IM_ENV_MARKERS = ("WEIXIN_TOKEN", "WECOM_BOT_ID", "DINGTALK_CLIENT_ID", "QQ_APP_ID")
# Group a one-blob ASR transcript into readable paragraphs at sentence boundaries —
# inserts paragraph breaks only, never alters a character (keeps 原话原序照搬).
_SENTENCE_BOUNDARY_RE = re.compile(r"(?<=[。！？!?…])")
_VERBATIM_PARAGRAPH_MAX_CHARS = 180
_VERBATIM_PARAGRAPH_MAX_SENTENCES = 4
_TRANSCRIPT_DOC_INSTRUCTION = (
    "用户要完整文案/逐字稿/转写全文/下载文案/原话时：直接在回复里单独一行输出 "
    "transcript_doc_media_tag 的值，把这份逐字稿文件发给用户——正文已是转写原文逐字照搬。"
    "不要再调用 doc_file_write，不要把逐字稿重排成结构化文档（加标题/小节、表格、要点都不行）。"
)


def _doc_file_delivery_available() -> bool:
    """True only when python-docx is importable AND this is a non-Feishu IM entry —
    the same gate apexnodes-doc-file-write uses, so the verbatim .docx export only
    surfaces where a Feishu doc link can't be opened."""
    import importlib.util

    try:
        if importlib.util.find_spec("docx") is None:
            return False
    except (ImportError, ValueError):
        return False
    return any((os.getenv(marker) or "").strip() for marker in _NON_FEISHU_IM_ENV_MARKERS)


def _document_cache_dir() -> Path:
    """Runtime document cache dir (a media-delivery safe root), matching doc_file_write."""
    home: str | None = None
    try:  # match the runtime's canonical HERMES_HOME resolution when available
        from hermes_constants import get_hermes_home

        home = str(get_hermes_home())
    except Exception:  # noqa: BLE001
        home = os.getenv("HERMES_HOME") or str(Path.home() / ".hermes")
    cache_dir = Path(home) / "cache" / "documents"
    cache_dir.mkdir(parents=True, exist_ok=True)
    return cache_dir


def _safe_filename(title: str) -> str:
    base = re.sub(r"[\\/\x00-\x1f]", " ", title or "").strip()
    base = re.sub(r"\s+", " ", base) or "视频逐字稿"
    return base[:60]


def _set_cjk_font(doc, font_name: str = "Microsoft YaHei") -> None:
    """Request a CJK font on the Normal style so viewers render Chinese (best-effort)."""
    try:
        from docx.oxml.ns import qn

        style = doc.styles["Normal"]
        style.font.name = font_name
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), font_name)
    except Exception:  # noqa: BLE001
        pass


def _verbatim_paragraphs(transcript: str) -> list[str]:
    """Split the transcript into readable paragraphs without changing any words.

    Respects segmentation the transcript already carries (newlines); otherwise groups
    sentences from a single ASR blob. Only paragraph breaks are added — the speaker's
    words and their order are preserved verbatim (the only processing the逐字稿 allows)."""
    text = (transcript or "").strip()
    if not text:
        return []
    explicit = [seg.strip() for seg in re.split(r"\n+", text) if seg.strip()]
    if len(explicit) > 1:
        return explicit
    paragraphs: list[str] = []
    buf: list[str] = []
    for sentence in (s for s in _SENTENCE_BOUNDARY_RE.split(text) if s.strip()):
        buf.append(sentence)
        if len(buf) >= _VERBATIM_PARAGRAPH_MAX_SENTENCES or sum(len(s) for s in buf) >= _VERBATIM_PARAGRAPH_MAX_CHARS:
            paragraphs.append("".join(buf))
            buf = []
    if buf:
        paragraphs.append("".join(buf))
    return paragraphs or [text]


def render_verbatim_transcript_docx(transcript: str, title: str) -> str | None:
    """Render the verbatim transcript to a .docx (heading + raw paragraphs) and return
    its path, or None on any failure. No markdown parsing and no headings/tables/bullets
    inside the body — the document is the speaker's original words, copied verbatim."""
    try:
        from docx import Document
    except ImportError:
        return None
    paragraphs = _verbatim_paragraphs(transcript)
    if not paragraphs:
        return None
    try:
        heading = (str(title or "").strip() or "视频逐字稿")[:120]
        doc = Document()
        _set_cjk_font(doc)
        doc.add_heading(heading, level=0)
        for para in paragraphs:
            doc.add_paragraph(para)
        dest = _document_cache_dir() / f"{_safe_filename(heading)}_逐字稿_{uuid.uuid4().hex[:8]}.docx"
        doc.save(str(dest))
        return str(dest)
    except Exception:  # noqa: BLE001 — a render failure must never fail the transcribe tool
        return None


def _handle_media_transcribe(args: dict, **_kwargs) -> str:
    if not isinstance(args, dict):
        return tool_error("media_transcribe expects a JSON object argument")
    video_path = str(args.get("video_path") or "").strip() or None
    url = str(args.get("url") or args.get("share_url") or "").strip() or None
    if not video_path and not url:
        return tool_error("请提供 video_path 或视频分享链接")
    if _use_gateway():
        return _gateway_media_transcribe(video_path, url)
    payload: dict[str, Any] = {"video_path": video_path, "url": url}
    if _routed_intent(args):
        payload["routed_intent"] = _routed_intent(args)
    try:
        result = _post("/media/transcribe", payload, timeout=_MEDIA_TRANSCRIBE_TIMEOUT_SECONDS)
    except RuntimeError as exc:
        return tool_error(f"视频转写失败: {exc}")
    # hc-341: pre-render the verbatim transcript so "完整文案/逐字稿/下载文案" is delivered
    # as the original speech, copied verbatim — the agent only emits the MEDIA line.
    transcript = str(result.get("transcript") or "").strip()
    if transcript and _doc_file_delivery_available():
        doc_path = render_verbatim_transcript_docx(transcript, str(result.get("title") or ""))
        if doc_path:
            result["transcript_doc_path"] = doc_path
            result["transcript_doc_media_tag"] = f"MEDIA:{doc_path}"
            result["transcript_doc_instruction"] = _TRANSCRIPT_DOC_INSTRUCTION
    return tool_result(**result)


def _handle_image_ocr(args: dict, **_kwargs) -> str:
    if not isinstance(args, dict):
        return tool_error("image_ocr expects a JSON object argument")
    raw_images = args.get("image_urls")
    image_urls = [str(item).strip() for item in raw_images if str(item).strip()] if isinstance(raw_images, list) else None
    payload = {
        "url": str(args.get("url") or args.get("share_url") or "").strip() or None,
        "image_urls": image_urls or None,
        "prompt": str(args.get("prompt") or args.get("question") or "").strip() or None,
    }
    if not payload["url"] and not payload["image_urls"]:
        return tool_error("请提供社媒图文链接 url 或图片地址 image_urls")
    if _use_gateway():
        return _gateway_image_ocr(payload)
    try:
        return tool_result(**_post("/media/image-understand", payload))
    except RuntimeError as exc:
        return tool_error(f"图片 OCR 失败: {exc}")


SOCIAL_BATCH_SUBMIT_SCHEMA = {
    "name": "social_batch_submit",
    "description": (
        "批量下载并转写一批社媒视频（抖音/快手/B站/小红书=下载+转写，YouTube=官方字幕，TikTok复用抖音）。"
        "适用于「把某作者点赞超10万的作品批量转写」「把这个抖音合集的视频都转成文字/表格」这类需求。"
        "传 urls（链接清单）或 creator_url（作者主页链接=枚举其全部作品，或某个抖音合集的分享链接=只枚举该合集；配 min_likes/top 自动筛选）。"
        "★这是异步任务：本工具只负责提交，返回 job_id；之后用 social_batch_status 轮询进度。"
        "平台会在 master 侧完成全部下载/转写/出表，你只需提交、轮询、把最终产物发给用户——"
        "绝不要自己写脚本下载、起进程、删文件。单条出 Word、多条出 Excel、飞书多条且给了 bitable_url 则写入多维表格。"
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "urls": {
                "type": "array",
                "items": {"type": "string"},
                "description": "视频链接清单（每个元素一条分享链接/口令）。与 creator_url 二选一。",
            },
            "creator_url": {
                "type": "string",
                "description": (
                    "作者主页链接（枚举其全部作品），或某个抖音合集的分享链接（douyin.com/collection/… 或 mix/…，"
                    "只枚举该合集内的作品）；平台会枚举并按 min_likes/top 筛选。与 urls 二选一。"
                    "注意：合集请发该【合集】自己的分享链接，仅凭作者主页+合集名无法定位某个合集。"
                ),
            },
            "platform": {"type": "string", "description": "可选，creator_url 的平台（douyin/kuaishou/bilibili/xiaohongshu/youtube/tiktok）。"},
            "min_likes": {"type": "integer", "description": "可选，creator_url 时只保留点赞≥该值的作品（如 100000）。"},
            "top": {"type": "integer", "description": "可选，最多取前 N 条（按点赞排序）。"},
            "bitable_url": {"type": "string", "description": "可选，飞书多条交付时写入的目标多维表格 /base/ 链接。"},
            "delivery_format": {
                "type": "string",
                "description": "可选，产物格式覆盖：xlsx=强制出 Excel 表格（用户明确要「表格/Excel」时，即使只有一条链接也出表）。缺省按条数自动（单条 Word、多条 Excel）。",
            },
        },
    },
}

SOCIAL_BATCH_STATUS_SCHEMA = {
    "name": "social_batch_status",
    "description": (
        "查询批量转写任务的进度与产物（配合 social_batch_submit 使用）。传 job_id。"
        "状态为 completed/partial 时，结果里 product 含产物：file 类（Word/Excel）给 download_url，"
        "飞书多维表格给 bitable_url——把它发给用户即可。仍在 running 时稍后再轮询，不要重复提交。"
    ),
    "parameters": {
        "type": "object",
        "properties": {"job_id": {"type": "string", "description": "social_batch_submit 返回的任务 id。"}},
        "required": ["job_id"],
    },
}


def _session_target() -> dict | None:
    """hc-371: capture the originating IM session so the master can push the batch
    result back here instead of the agent polling. The v0.17 gateway mirrors the
    turn's session onto ``HERMES_SESSION_*`` (ContextVar → os.environ → tool thread);
    ``get_session_env`` reads it, with a plain ``os.getenv`` fallback if the gateway
    helper isn't importable. ``None`` (no chat_id) ⇒ master can't route ⇒ poll."""
    try:
        from gateway.session_context import get_session_env  # type: ignore
    except Exception:  # pragma: no cover - exercised inside the v0.17 agent only
        def get_session_env(name: str, default: str = "") -> str:
            return os.getenv(name, default)

    chat_id = (get_session_env("HERMES_SESSION_CHAT_ID", "") or "").strip()
    if not chat_id:
        return None
    target = {"chat_id": chat_id}
    for key, env in (("platform", "HERMES_SESSION_PLATFORM"), ("user_id", "HERMES_SESSION_USER_ID")):
        value = (get_session_env(env, "") or "").strip()
        if value:
            target[key] = value
    return target


def _decorate_batch_submit(result: dict[str, Any]) -> dict[str, Any]:
    if result.get("async_delivery"):
        # The master will push the finished result straight to this chat — no poll.
        result["_instruction"] = (
            "已提交批量任务。任务在 master 侧异步执行，完成后平台会【自动把结果推送到当前会话】，"
            "你无需轮询、可以直接结束本轮；完成时用户会收到总结与产物链接。"
            "请勿重复提交、勿自己下载或起脚本。"
        )
    else:
        result["_instruction"] = (
            "已提交批量任务。请记住 job_id，稍后用 social_batch_status 轮询；任务在 master 侧异步执行，"
            "完成前不要重复提交、不要自己下载或起脚本。完成后把产物链接发给用户。"
        )
    return result


def _decorate_batch_status(result: dict[str, Any]) -> dict[str, Any]:
    status = result.get("status")
    if status in ("completed", "partial"):
        result["_instruction"] = (
            "任务已完成。result.note 是给用户的总结；product 是产物："
            "若有 download_url(Word/Excel) 就把文件发给用户，若有 bitable_url 就把多维表格链接发给用户。"
        )
    elif status == "failed":
        result["_instruction"] = "任务失败，把 result.note / error 里的原因如实告诉用户。"
    else:
        result["_instruction"] = "任务仍在进行中（still running）。稍等片刻再用 social_batch_status 轮询，不要重复提交。"
    return result


def _handle_social_batch_submit(args: dict, **_kwargs) -> str:
    if not isinstance(args, dict):
        return tool_error("social_batch_submit expects a JSON object argument")
    urls = args.get("urls")
    if isinstance(urls, str):
        urls = [urls]
    payload = {
        "urls": [str(u).strip() for u in urls if str(u).strip()] if isinstance(urls, list) else None,
        "creator_url": str(args.get("creator_url") or "").strip() or None,
        "platform": str(args.get("platform") or "").strip() or None,
        "min_likes": int(args.get("min_likes") or 0),
        "top": args.get("top"),
        "bitable_url": str(args.get("bitable_url") or "").strip() or None,
        # hc-450 PR2: explicit product-format override(今日仅 "xlsx")——legacy
        # /media/batch/submit 与网关 /tools/v1/social/batch/submit 两端点同形接收。
        "delivery_format": str(args.get("delivery_format") or "").strip().lower() or None,
        # hc-371: routing snapshot; the master stores it and pushes the result on
        # completion. Always captured (harmless when the feature is off — nothing
        # reads it until the master's async-delivery flag is armed for this agent;
        # 桌面形态没有 IM 会话 env ⇒ None ⇒ 轮询模式).
        "delivery_target": _session_target(),
    }
    # hc-450 PR1 观测透传:legacy 端点打 intent_router_hit 结构化行;网关端点当前
    # 忽略该字段(服务端模型未收,pydantic 默认丢弃)——共用同一 payload,服务端补字段后即生效。
    if _routed_intent(args):
        payload["routed_intent"] = _routed_intent(args)
    if not payload["urls"] and not payload["creator_url"]:
        return tool_error("请提供 urls(链接清单) 或 creator_url(作者主页链接，或某个抖音合集的分享链接)。")
    if _use_gateway():
        try:
            result = _gateway.unwrap(
                _gateway.request_json("POST", "/tools/v1/social/batch/submit", payload, timeout=120)
            )
        except _gateway.GatewayError as exc:
            return tool_error(f"批量任务提交失败: {exc}")
        return tool_result(**_decorate_batch_submit(result))
    try:
        result = _post("/media/batch/submit", payload, timeout=120)
    except RuntimeError as exc:
        return tool_error(f"批量任务提交失败: {exc}")
    return tool_result(**_decorate_batch_submit(result))


def _handle_social_batch_status(args: dict, **_kwargs) -> str:
    if not isinstance(args, dict):
        return tool_error("social_batch_status expects a JSON object argument")
    job_id = str(args.get("job_id") or "").strip()
    if not job_id:
        return tool_error("请提供 job_id。")
    if _use_gateway():
        try:
            result = _gateway.unwrap(
                _gateway.request_json("GET", f"/tools/v1/social/batch/status/{job_id}", timeout=60)
            )
        except _gateway.GatewayError as exc:
            return tool_error(f"查询批量任务失败: {exc}")
        return tool_result(**_decorate_batch_status(result))
    try:
        result = _get(f"/media/batch/status/{job_id}")
    except RuntimeError as exc:
        return tool_error(f"查询批量任务失败: {exc}")
    return tool_result(**_decorate_batch_status(result))


def register(ctx):
    ctx.register_tool(
        name="social_download",
        toolset="skills",
        schema=SOCIAL_DOWNLOAD_SCHEMA,
        handler=_handle_social_download,
        check_fn=_check,
        requires_env=[],
        description=SOCIAL_DOWNLOAD_SCHEMA["description"],
        emoji="⬇️",
    )
    ctx.register_tool(
        name="media_transcribe",
        toolset="skills",
        schema=MEDIA_TRANSCRIBE_SCHEMA,
        handler=_handle_media_transcribe,
        check_fn=_check,
        requires_env=[],
        description=MEDIA_TRANSCRIBE_SCHEMA["description"],
        emoji="📝",
    )
    ctx.register_tool(
        name="image_ocr",
        toolset="skills",
        schema=IMAGE_OCR_SCHEMA,
        handler=_handle_image_ocr,
        check_fn=_check_image,
        requires_env=[],
        description=IMAGE_OCR_SCHEMA["description"],
        emoji="🖼️",
    )
    ctx.register_tool(
        name="social_batch_submit",
        toolset="skills",
        schema=SOCIAL_BATCH_SUBMIT_SCHEMA,
        handler=_handle_social_batch_submit,
        check_fn=_check,
        requires_env=[],
        description=SOCIAL_BATCH_SUBMIT_SCHEMA["description"],
        emoji="📦",
    )
    ctx.register_tool(
        name="social_batch_status",
        toolset="skills",
        schema=SOCIAL_BATCH_STATUS_SCHEMA,
        handler=_handle_social_batch_status,
        check_fn=_check,
        requires_env=[],
        description=SOCIAL_BATCH_STATUS_SCHEMA["description"],
        emoji="🔎",
    )
