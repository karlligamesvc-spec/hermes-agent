"""ApexNodes document-file write tool (hc-215).

Non-Feishu IM entries (WeChat / QQ / WeCom / DingTalk) cannot open a Feishu doc
link, so the hc-164 long-output policy's ``feishu_doc_write`` delivery is wrong
for them: without owner Feishu credentials it fails (and the agent falls back to
dumping the whole text into chat), and *with* credentials it hands the user a
Feishu URL they can't open. This tool renders the full markdown reply to a
``.docx`` file under the runtime's document cache (a media-delivery safe root)
and returns the path; the agent then emits ``MEDIA:<path>`` so the in-container
gateway delivers the file natively (WeChat/QQ ``send_document``). The owner never
needs any third-party doc authorization and the user sees the file instantly.

Availability (``check_fn``) is gated on ``python-docx`` being importable (baked
into the runtime image / bundled with the desktop install) AND the delivery
surface being a non-Feishu one where a ``MEDIA:`` file is the right output:
  * a non-Feishu IM credential marker in the container env
    (``WEIXIN_TOKEN`` / ``WECOM_BOT_ID`` / ``DINGTALK_CLIENT_ID`` / ``QQ_APP_ID``), or
  * the desktop app itself (``HERMES_DESKTOP=1``, set by the electron shell on the
    dashboard backend it spawns; hc-565) — the desktop is its own non-Feishu
    delivery surface that renders ``MEDIA:`` files natively.

so the tool stays invisible to Feishu and managed cloud agents (which keep
``feishu_doc_write``) and surfaces on the IM entries and the desktop that need
it. The long-output policy prefers this tool over ``feishu_doc_write`` whenever
it is available, and still falls back to chat full-text on any failure.
"""

from __future__ import annotations

import os
import re
import uuid
from pathlib import Path

from tools.registry import tool_error, tool_result

# Container env markers for the non-Feishu IM entries this tool serves. Injected
# per entry by the platform (docker_manager._build_env). Kept in sync with the
# entry credentials documented in hc-188/hc-190/hc-202.
_NON_FEISHU_IM_ENV_MARKERS = (
    "WEIXIN_TOKEN",       # 个人微信 (iLink)
    "WECOM_BOT_ID",       # 企业微信
    "DINGTALK_CLIENT_ID",  # 钉钉
    "QQ_APP_ID",          # QQ 机器人
)

# hc-565 desktop entry: the electron shell spawns the dashboard backend with
# HERMES_DESKTOP=1 (apps/desktop/electron/main.cjs). The desktop app is a
# non-Feishu delivery surface that renders MEDIA: files natively, so doc_file_write
# must surface there too — otherwise the tool's own IM-env gate would hide it on
# desktop (audit hc-561 ①#22). Cloud containers never set this, so cloud behavior
# is unchanged.
_DESKTOP_ENV_MARKER = "HERMES_DESKTOP"

DOC_FILE_WRITE_SCHEMA = {
    "name": "doc_file_write",
    "description": (
        "Render a long structured reply (the full markdown text) to a .docx "
        "file and return its local path, for delivery on non-Feishu IM entries "
        "(WeChat/QQ/WeCom/DingTalk) where a Feishu doc link cannot be opened. "
        "After calling this, put a line `MEDIA:<file_path>` in your reply so the "
        "file is sent to the user, plus a short (<=5 line) summary. Markdown "
        "headings/lists/tables/code/bold are rendered as native docx structure."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "content": {
                "type": "string",
                "description": "Full markdown (or plain) text to render into the .docx file.",
            },
            "title": {
                "type": "string",
                "description": "Document title / file name (defaults to a generic report name).",
            },
        },
        "required": ["content"],
    },
}


def _check_doc_file_write() -> bool:
    """Available when python-docx is present AND the surface is non-Feishu: a
    non-Feishu IM entry (cloud/desktop) OR the desktop app itself (hc-565)."""
    import importlib.util

    try:
        if importlib.util.find_spec("docx") is None:
            return False
    except (ImportError, ValueError):
        return False
    if (os.getenv(_DESKTOP_ENV_MARKER) or "").strip():
        return True
    return any((os.getenv(marker) or "").strip() for marker in _NON_FEISHU_IM_ENV_MARKERS)


def _document_cache_dir() -> Path:
    """Return the runtime document cache dir (a media-delivery safe root)."""
    home: str | None = None
    try:  # match the runtime's canonical HERMES_HOME resolution when available
        from hermes_constants import get_hermes_home

        home = str(get_hermes_home())
    except Exception:  # noqa: BLE001
        home = os.getenv("HERMES_HOME") or str(Path.home() / ".hermes")
    cache_dir = Path(home) / "cache" / "documents"
    cache_dir.mkdir(parents=True, exist_ok=True)
    return cache_dir


def _safe_filename(title: str) -> str:
    base = re.sub(r"[\\/\x00-\x1f]", " ", title or "").strip()
    base = re.sub(r"\s+", " ", base) or "Agent report"
    return base[:60]


# ── markdown → docx ─────────────────────────────────────────────────────────

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_DIVIDER_RE = re.compile(r"^(-{3,}|\*{3,}|_{3,})$")
_BULLET_RE = re.compile(r"^\s*[-*+]\s+(.+)$")
_ORDERED_RE = re.compile(r"^\s*\d+[.)]\s+(.+)$")
_TABLE_SEP_RE = re.compile(r"^\|?[\s:|-]+\|?$")
_INLINE_TOKEN_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)", re.DOTALL)


def _set_cjk_font(doc, font_name: str = "Microsoft YaHei") -> None:
    """Make the Normal style request a CJK font so viewers render Chinese.

    Best-effort: even without it, the WeChat/QQ document viewer substitutes a
    CJK font for missing glyphs, so failure here must never break rendering.
    """
    try:
        from docx.oxml.ns import qn

        style = doc.styles["Normal"]
        style.font.name = font_name
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), font_name)
    except Exception:  # noqa: BLE001
        pass


def _add_inline_runs(paragraph, text: str) -> None:
    """Add **bold** / `code` aware runs to a paragraph."""
    for part in _INLINE_TOKEN_RE.split(text or ""):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def _split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def render_markdown_to_docx(content: str, title: str, dest: Path) -> Path:
    """Render markdown ``content`` to a .docx at ``dest``. Returns ``dest``.

    Pure rendering — separated from the handler so it is unit-testable. Raises on
    a genuine python-docx failure so the handler can surface a tool_error and the
    agent falls back to chat full-text.
    """
    from docx import Document

    doc = Document()
    _set_cjk_font(doc)
    if title:
        doc.add_heading(title[:120], level=0)

    lines = (content or "").split("\n")
    i, total = 0, len(lines)
    while i < total:
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):  # fenced code block
            buf: list[str] = []
            i += 1
            while i < total and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # skip closing fence when present
            para = doc.add_paragraph()
            run = para.add_run("\n".join(buf))
            run.font.name = "Consolas"
            continue

        if (
            stripped.startswith("|")
            and "|" in stripped[1:]
            and i + 1 < total
            and _TABLE_SEP_RE.match(lines[i + 1].strip())
            and "-" in lines[i + 1]
        ):
            rows = [_split_table_row(stripped)]
            i += 2  # header + separator
            while i < total and lines[i].strip().startswith("|") and "|" in lines[i].strip()[1:]:
                rows.append(_split_table_row(lines[i].strip()))
                i += 1
            columns = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=columns)
            try:
                table.style = "Table Grid"
            except Exception:  # noqa: BLE001 — style availability varies by template
                pass
            for r, row in enumerate(rows):
                cells = row + [""] * (columns - len(row))
                for c, cell_text in enumerate(cells):
                    cell_para = table.cell(r, c).paragraphs[0]
                    _add_inline_runs(cell_para, cell_text)
                    if r == 0:  # header row bold
                        for run in cell_para.runs:
                            run.bold = True
            continue

        if _DIVIDER_RE.match(stripped):
            doc.add_paragraph().add_run("―" * 20)
            i += 1
            continue

        heading = _HEADING_RE.match(stripped)
        if heading:
            level = min(len(heading.group(1)), 6)
            doc.add_heading(heading.group(2).strip(), level=level)
            i += 1
            continue

        ordered = _ORDERED_RE.match(line)
        if ordered:
            _add_inline_runs(doc.add_paragraph(style="List Number"), ordered.group(1).strip())
            i += 1
            continue

        bullet = _BULLET_RE.match(line)
        if bullet:
            _add_inline_runs(doc.add_paragraph(style="List Bullet"), bullet.group(1).strip())
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        _add_inline_runs(doc.add_paragraph(), stripped)
        i += 1

    doc.save(str(dest))
    return dest


def _handle_doc_file_write(args: dict, **_kwargs) -> str:
    if not isinstance(args, dict):
        return tool_error("doc_file_write expects a JSON object argument")

    content = args.get("content")
    if content is None or (isinstance(content, str) and not content.strip()):
        return tool_error("content is required")

    title = str(args.get("title") or "Agent report").strip() or "Agent report"

    try:
        cache_dir = _document_cache_dir()
        dest = cache_dir / f"{_safe_filename(title)}_{uuid.uuid4().hex[:8]}.docx"
        render_markdown_to_docx(str(content), title, dest)
    except ImportError:
        return tool_error("python-docx is not installed in this runtime image")
    except Exception as exc:  # noqa: BLE001 — surface so the agent falls back to chat full-text
        return tool_error(f"Failed to render the document file: {exc}")

    path = str(dest)
    return tool_result(
        success=True,
        file_path=path,
        media_tag=f"MEDIA:{path}",
        instruction=(
            "在回复里单独一行输出 media_tag 的值(例如 MEDIA:/path/to/file.docx)即可把文件发给用户;"
            "另附不超过五行的核心摘要。不要发任何飞书或外部链接。"
        ),
    )


def register(ctx):
    ctx.register_tool(
        name="doc_file_write",
        toolset="doc_delivery",
        schema=DOC_FILE_WRITE_SCHEMA,
        handler=_handle_doc_file_write,
        check_fn=_check_doc_file_write,
        requires_env=[],
        description=DOC_FILE_WRITE_SCHEMA["description"],
        emoji="📄",
    )
